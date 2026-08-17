from dataclasses import dataclass
from io import BytesIO
from zipfile import BadZipFile, ZipFile

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError
from docx.table import Table
from docx.text.paragraph import Paragraph
from lxml.etree import XMLSyntaxError
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from settings import settings

TEXT_CONTENT_TYPE = "text/plain"
PDF_CONTENT_TYPE = "application/pdf"
DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument."
    "wordprocessingml.document"
)

SUPPORTED_DOCUMENT_CONTENT_TYPES = frozenset(
    {
        TEXT_CONTENT_TYPE,
        PDF_CONTENT_TYPE,
        DOCX_CONTENT_TYPE,
    },
)


class DocumentTextExtractionError(Exception):
    message: str

    def __init__(self) -> None:
        super().__init__(self.message)


class UnsupportedDocumentFormatError(DocumentTextExtractionError):
    message = "Document content type is not supported."


class InvalidDocumentFormatError(DocumentTextExtractionError):
    message = "Document format is invalid."


class PasswordProtectedPdfError(DocumentTextExtractionError):
    message = "Password-protected PDF is not supported."


class DocumentExtractionLimitError(DocumentTextExtractionError):
    message = "Document exceeds processing limits."


class DocumentTextEmptyError(DocumentTextExtractionError):
    message = "Document contains no extractable text."


@dataclass(frozen=True, slots=True)
class ExtractedBlock:
    text: str
    page_number: int | None


@dataclass(frozen=True, slots=True)
class ExtractedDocument:
    blocks: tuple[ExtractedBlock, ...]

    @property
    def text(self) -> str:
        return "\n\n".join(
            block.text
            for block in self.blocks
        )


def extract_document_text(
    *,
    content: bytes,
    content_type: str,
) -> ExtractedDocument:
    if content_type == TEXT_CONTENT_TYPE:
        return _extract_plain_text(content)

    if content_type == PDF_CONTENT_TYPE:
        return _extract_pdf(content)

    if content_type == DOCX_CONTENT_TYPE:
        return _extract_docx(content)

    raise UnsupportedDocumentFormatError()


def validate_document_upload_content(
    *,
    content: bytes,
    content_type: str,
) -> None:
    if content_type == TEXT_CONTENT_TYPE:
        try:
            content.decode("utf-8")
        except UnicodeDecodeError as exc:
            raise InvalidDocumentFormatError() from exc

        return

    if content_type == PDF_CONTENT_TYPE:
        if not content.startswith(b"%PDF-"):
            raise InvalidDocumentFormatError()

        return

    if content_type == DOCX_CONTENT_TYPE:
        _validate_docx_archive(content)
        return

    raise UnsupportedDocumentFormatError()


def _extract_plain_text(
    content: bytes,
) -> ExtractedDocument:
    try:
        text = content.decode("utf-8")
    except UnicodeDecodeError as exc:
        raise InvalidDocumentFormatError() from exc

    blocks: list[ExtractedBlock] = []
    _append_block(
        blocks=blocks,
        text=text,
        page_number=None,
        extracted_chars=0,
    )

    return _build_extracted_document(blocks)


def _extract_pdf(
    content: bytes,
) -> ExtractedDocument:
    try:
        reader = PdfReader(
            BytesIO(content),
            strict=True,
        )

        if reader.is_encrypted:
            raise PasswordProtectedPdfError()

        if len(reader.pages) > settings.document_pdf_max_pages:
            raise DocumentExtractionLimitError()

        blocks: list[ExtractedBlock] = []
        extracted_chars = 0

        for page_number, page in enumerate(
            reader.pages,
            start=1,
        ):
            contents = page.get_contents()

            if (
                contents is not None
                and len(contents.get_data())
                > settings.document_pdf_max_page_content_bytes
            ):
                raise DocumentExtractionLimitError()

            extracted_chars = _append_block(
                blocks=blocks,
                text=page.extract_text() or "",
                page_number=page_number,
                extracted_chars=extracted_chars,
            )

        return _build_extracted_document(blocks)
    except DocumentTextExtractionError:
        raise
    except (
        PdfReadError,
        IndexError,
        KeyError,
        OSError,
        TypeError,
        ValueError,
    ) as exc:
        raise InvalidDocumentFormatError() from exc


def _extract_docx(
    content: bytes,
) -> ExtractedDocument:
    _validate_docx_archive(content)

    try:
        document = DocxDocument(BytesIO(content))
    except (
        PackageNotFoundError,
        XMLSyntaxError,
        KeyError,
        OSError,
        TypeError,
        ValueError,
    ) as exc:
        raise InvalidDocumentFormatError() from exc

    blocks: list[ExtractedBlock] = []
    extracted_chars = 0
    paragraph_count = 0

    for element in document.iter_inner_content():
        if isinstance(element, Paragraph):
            paragraph_count += 1
            _ensure_docx_paragraph_limit(paragraph_count)

            extracted_chars = _append_block(
                blocks=blocks,
                text=element.text,
                page_number=None,
                extracted_chars=extracted_chars,
            )
            continue

        if isinstance(element, Table):
            for row in element.rows:
                row_cells: list[str] = []

                for cell in row.cells:
                    paragraph_count += len(cell.paragraphs)
                    _ensure_docx_paragraph_limit(paragraph_count)

                    cell_text = "\n".join(
                        paragraph.text
                        for paragraph in cell.paragraphs
                    ).strip()

                    if cell_text:
                        row_cells.append(cell_text)

                extracted_chars = _append_block(
                    blocks=blocks,
                    text=" | ".join(row_cells),
                    page_number=None,
                    extracted_chars=extracted_chars,
                )

    return _build_extracted_document(blocks)


def _validate_docx_archive(
    content: bytes,
) -> None:
    try:
        with ZipFile(BytesIO(content)) as archive:
            infos = archive.infolist()
    except (
        BadZipFile,
        OSError,
        ValueError,
    ) as exc:
        raise InvalidDocumentFormatError() from exc

    names = {
        info.filename.lower()
        for info in infos
    }
    required_names = {
        "[content_types].xml",
        "word/document.xml",
    }

    if not required_names.issubset(names):
        raise InvalidDocumentFormatError()

    if "word/vbaproject.bin" in names:
        raise InvalidDocumentFormatError()

    if any(info.flag_bits & 0x1 for info in infos):
        raise InvalidDocumentFormatError()

    total_uncompressed_bytes = sum(
        info.file_size
        for info in infos
    )

    if (
        total_uncompressed_bytes
        > settings.document_docx_max_uncompressed_bytes
    ):
        raise DocumentExtractionLimitError()


def _ensure_docx_paragraph_limit(
    paragraph_count: int,
) -> None:
    if paragraph_count > settings.document_docx_max_paragraphs:
        raise DocumentExtractionLimitError()


def _append_block(
    *,
    blocks: list[ExtractedBlock],
    text: str,
    page_number: int | None,
    extracted_chars: int,
) -> int:
    normalized_text = "\n".join(
        line.strip()
        for line in text.splitlines()
        if line.strip()
    )

    if not normalized_text:
        return extracted_chars

    separator_size = 2 if blocks else 0
    new_extracted_chars = (
        extracted_chars
        + separator_size
        + len(normalized_text)
    )

    if (
        new_extracted_chars
        > settings.document_extracted_text_max_chars
    ):
        raise DocumentExtractionLimitError()

    blocks.append(
        ExtractedBlock(
            text=normalized_text,
            page_number=page_number,
        ),
    )

    return new_extracted_chars


def _build_extracted_document(
    blocks: list[ExtractedBlock],
) -> ExtractedDocument:
    if not blocks:
        raise DocumentTextEmptyError()

    return ExtractedDocument(
        blocks=tuple(blocks),
    )